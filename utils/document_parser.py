import PyPDF2
import io
import docx
from PyPDF2 import PdfReader

def parse_pdf(file_stream):
    """Parses text from a PDF file stream."""
    text = ""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Could not read PDF file: {e}")
    return text

def parse_docx(file_stream):
    """Parses text from a DOCX file stream."""
    text = ""
    try:
        doc = docx.Document(file_stream)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {e}")
    return text

def parse_document(uploaded_file):
    """Parses an uploaded file (PDF or DOCX) and returns its text content."""
    file_stream = io.BytesIO(uploaded_file.getvalue())
    file_type = uploaded_file.type

    if "pdf" in file_type:
        return parse_pdf(file_stream)
    elif "vnd.openxmlformats-officedocument.wordprocessingml.document" in file_type:
        return parse_docx(file_stream)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def parse_document(file_stream, file_type: str) -> str:
    """Parse uploaded PDF or DOCX file and return extracted text."""
    text = ""

    if file_type == "application/pdf":
        reader = PdfReader(file_stream)
        for page in reader.pages:
            text += page.extract_text() + "\n"

    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file_stream)
        for para in doc.paragraphs:
            text += para.text + "\n"

    else:
        raise ValueError("Unsupported file type")

    return text.strip()